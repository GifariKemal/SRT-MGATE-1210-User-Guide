#!/usr/bin/env python3
"""
Convert USER_GUIDE.md to DOCX with proper image sizing and page numbers.
Author: Claude AI for SURIOTA
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Paths
SCRIPT_DIR = Path(__file__).parent
PROJECT_DIR = SCRIPT_DIR.parent
MD_FILE = PROJECT_DIR / "docs" / "USER_GUIDE.md"
SCREENSHOTS_DIR = PROJECT_DIR / "assets" / "screenshots"
OUTPUT_DIR = PROJECT_DIR / "output"
OUTPUT_FILE = OUTPUT_DIR / "Panduan_Pengguna_Gateway_Config_App.docx"

# Document settings
IMAGE_WIDTH = Inches(1.8)  # Smaller size for side-by-side display
PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(1)

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

# Global counters for captions
figure_counter = {}  # Track figure numbers per section
table_counter = {}   # Track table numbers per section
current_section = [0]  # Current section number (using list for mutability)

# Color definitions for callout boxes (using hex strings directly)
COLORS = {
    'tip': {'bg': 'E8F5E9', 'border': '4CAF50', 'icon': '💡', 'title_color': RGBColor(76, 175, 80)},      # Green
    'warning': {'bg': 'FFF3E0', 'border': 'FF9800', 'icon': '⚠️', 'title_color': RGBColor(255, 152, 0)},  # Orange
    'note': {'bg': 'E3F2FD', 'border': '2196F3', 'icon': 'ℹ️', 'title_color': RGBColor(33, 150, 243)},    # Blue
    'danger': {'bg': 'FFEBEE', 'border': 'F44336', 'icon': '❌', 'title_color': RGBColor(244, 67, 54)},   # Red
}

def add_callout_box(doc, box_type, title, content):
    """
    Add a styled callout/info box.
    box_type: 'tip', 'warning', 'note', 'danger'
    """
    colors = COLORS.get(box_type, COLORS['note'])

    # Create a table with 1 row and 1 cell for the box
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table width to 90% of page
    table.autofit = False
    for cell in table.rows[0].cells:
        cell.width = Inches(6)

    cell = table.rows[0].cells[0]

    # Set cell shading (background color)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), colors['bg'])
    cell._tc.get_or_add_tcPr().append(shading)

    # Set cell border
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')  # Border width
        border.set(qn('w:color'), colors['border'])
        tcBorders.append(border)

    tcPr.append(tcBorders)

    # Add title with icon
    p = cell.paragraphs[0]
    run = p.add_run(f"{colors['icon']} {title}")
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = colors['title_color']

    # Add content
    p2 = cell.add_paragraph()
    run2 = p2.add_run(content)
    run2.font.size = Pt(10)
    run2.font.name = 'Calibri'

    # Add spacing after
    doc.add_paragraph()

def add_watermark(doc, text="SURIOTA"):
    """Add a diagonal watermark to the document header."""
    section = doc.sections[0]
    header = section.header

    # Create watermark using a shape in header
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Note: python-docx doesn't directly support watermarks
    # We'll add it as semi-transparent text in header instead
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(220, 220, 220)  # Very light gray
    run.font.name = 'Calibri'

def add_version_history(doc):
    """Add version history table after cover page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RIWAYAT REVISI DOKUMEN")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()

    # Create version history table
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    headers = ['Versi', 'Tanggal', 'Penulis', 'Deskripsi Perubahan']
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        # Set header background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '006699')
        cell._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Version entries
    versions = [
        ('1.0.0', 'Desember 2025', 'Tim SURIOTA', 'Rilis awal dokumen'),
        ('', '', '', ''),
        ('', '', '', ''),
    ]

    for row_idx, version_data in enumerate(versions, start=1):
        for col_idx, text in enumerate(version_data):
            cell = table.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'

    doc.add_paragraph()
    doc.add_page_break()

def add_glossary(doc):
    """Add glossary/terminology section at the end."""
    doc.add_page_break()

    # Title
    doc.add_heading("Glosarium / Daftar Istilah", level=1)

    doc.add_paragraph()

    # Glossary table
    terms = [
        ('API', 'Application Programming Interface - antarmuka untuk komunikasi antar aplikasi'),
        ('BLE', 'Bluetooth Low Energy - protokol Bluetooth hemat energi'),
        ('Gateway', 'Perangkat penghubung antara sensor dan server/cloud'),
        ('HTTP', 'Hypertext Transfer Protocol - protokol komunikasi web'),
        ('IIoT', 'Industrial Internet of Things - IoT untuk industri'),
        ('JSON', 'JavaScript Object Notation - format pertukaran data'),
        ('Modbus RTU', 'Protokol komunikasi serial untuk perangkat industri'),
        ('Modbus TCP', 'Protokol Modbus melalui jaringan TCP/IP'),
        ('MQTT', 'Message Queuing Telemetry Transport - protokol messaging IoT'),
        ('Register', 'Alamat memori pada perangkat Modbus untuk menyimpan data'),
        ('RTU', 'Remote Terminal Unit - unit terminal jarak jauh'),
        ('Slave ID', 'Alamat unik perangkat Modbus dalam jaringan'),
        ('TCP/IP', 'Transmission Control Protocol/Internet Protocol'),
        ('TLS/SSL', 'Transport Layer Security - enkripsi komunikasi'),
    ]

    table = doc.add_table(rows=len(terms) + 1, cols=2)
    table.style = 'Table Grid'

    # Header
    header_cells = table.rows[0].cells
    for i, text in enumerate(['Istilah', 'Definisi']):
        p = header_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '006699')
        header_cells[i]._tc.get_or_add_tcPr().append(shading)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Set column widths
    header_cells[0].width = Inches(1.5)
    header_cells[1].width = Inches(5)

    # Terms
    for row_idx, (term, definition) in enumerate(terms, start=1):
        cells = table.rows[row_idx].cells
        cells[0].width = Inches(1.5)
        cells[1].width = Inches(5)

        p1 = cells[0].paragraphs[0]
        run1 = p1.add_run(term)
        run1.font.bold = True
        run1.font.size = Pt(10)
        run1.font.name = 'Calibri'

        p2 = cells[1].paragraphs[0]
        run2 = p2.add_run(definition)
        run2.font.size = Pt(10)
        run2.font.name = 'Calibri'

    doc.add_paragraph()

def add_page_number(run):
    """Add page number field to a run."""
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def setup_header_footer(doc):
    """Add header with document series and footer with author and page numbers."""
    section = doc.sections[0]

    # === HEADER ===
    header = section.header
    header.is_linked_to_previous = False

    # Header paragraph - Document Series
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("SRT-MGATE-1210 User Guide Series")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.font.bold = True

    # Add separator line
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Panduan Pengguna Gateway Config App v1.0")
    run2.font.size = Pt(9)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(128, 128, 128)
    run2.font.italic = True

    # === FOOTER ===
    footer = section.footer
    footer.is_linked_to_previous = False

    # Footer line 1 - Company info
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("PT Surya Inovasi Prioritas (SURIOTA)")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.font.bold = True

    # Footer line 2 - Page number
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run2 = p2.add_run("Halaman ")
    run2.font.size = Pt(9)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(100, 100, 100)

    # Add page number field
    add_page_number(p2.add_run())

    run3 = p2.add_run(" | www.suriota.com")
    run3.font.size = Pt(9)
    run3.font.name = "Calibri"
    run3.font.color.rgb = RGBColor(100, 100, 100)

def setup_styles(doc):
    """Setup document styles."""
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 102, 153)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 102, 153)

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 128, 128)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True

    # Normal text
    normal = styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)

    return styles

def add_cover_page(doc):
    """Add professional cover page."""
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PANDUAN PENGGUNA")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Gateway Config App")
    run.font.size = Pt(24)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 128, 128)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nAplikasi Konfigurasi untuk\nSRT-MGATE-1210 Modbus Gateway IIoT")
    run.font.size = Pt(14)
    run.font.name = "Calibri"

    # Add spacing
    for _ in range(8):
        doc.add_paragraph()

    # Version info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Versi Dokumen: 1.0.0")
    run.font.size = Pt(12)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Desember 2025")
    run.font.size = Pt(12)
    run.font.name = "Calibri"

    # Company info
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PT Surya Inovasi Prioritas (SURIOTA)")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 102, 153)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("www.suriota.com")
    run.font.size = Pt(11)
    run.font.name = "Calibri"

    # Page break after cover
    doc.add_page_break()

def add_toc_field(doc):
    """Add automatic Table of Contents field like Word with dots and hyperlinks."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    # Create TOC field with proper switches:
    # \o "1-3" = include heading levels 1-3
    # \h = hyperlinks
    # \z = hide page numbers in web view
    # \u = use paragraph outline level
    # \p = separator between entry and page number (dots)
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')

    # Add placeholder text that will be replaced when field is updated
    placeholder = create_element('w:t')
    placeholder.text = 'Right-click and select "Update Field" to generate table of contents'

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(placeholder)
    run._r.append(fldChar3)

def setup_toc_styles(doc):
    """Setup TOC styles with tab stops for dots leader."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Get or create TOC 1 style
    styles = doc.styles

    for toc_level in range(1, 4):
        style_name = f'TOC {toc_level}'
        try:
            toc_style = styles[style_name]
        except KeyError:
            # Style doesn't exist, skip
            continue

        # Set font
        toc_style.font.name = 'Calibri'
        toc_style.font.size = Pt(11) if toc_level == 1 else Pt(10)

        # Set paragraph format
        pf = toc_style.paragraph_format
        if toc_level == 1:
            pf.left_indent = Cm(0)
            toc_style.font.bold = True
        elif toc_level == 2:
            pf.left_indent = Cm(0.5)
        else:
            pf.left_indent = Cm(1.0)

        pf.space_before = Pt(3)
        pf.space_after = Pt(3)

def add_table_of_contents(doc):
    """Add table of contents page with proper Word-style formatting."""
    # Setup TOC styles
    setup_toc_styles(doc)

    # TOC Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAFTAR ISI")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()
    doc.add_paragraph()

    # Add automatic TOC field (will show dots and hyperlinks when updated in Word)
    add_toc_field(doc)

    doc.add_paragraph()
    doc.add_paragraph()

    # Instructions for user
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Untuk menampilkan Daftar Isi dengan hyperlink dan nomor halaman:")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Klik kanan pada teks di atas → pilih 'Update Field' → pilih 'Update entire table'")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.font.name = "Calibri"

    doc.add_page_break()

    # === DAFTAR GAMBAR ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAFTAR GAMBAR")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()

    # Add List of Figures field
    p = doc.add_paragraph()
    run = p.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = ' TOC \\h \\z \\c "Gambar" '  # List of figures with caption label "Gambar"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')

    placeholder = create_element('w:t')
    placeholder.text = 'Update Field untuk menampilkan daftar gambar'

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(placeholder)
    run._r.append(fldChar3)

    doc.add_paragraph()
    doc.add_page_break()

    # === DAFTAR TABEL ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DAFTAR TABEL")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 102, 153)

    doc.add_paragraph()

    # Add List of Tables field
    p = doc.add_paragraph()
    run = p.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = ' TOC \\h \\z \\c "Tabel" '  # List of tables with caption label "Tabel"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')

    placeholder = create_element('w:t')
    placeholder.text = 'Update Field untuk menampilkan daftar tabel'

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(placeholder)
    run._r.append(fldChar3)

    doc.add_paragraph()
    doc.add_page_break()

def find_image(image_name):
    """Find image file in screenshots directory."""
    # Clean the image name
    image_name = image_name.strip()

    # Try exact match first
    image_path = SCREENSHOTS_DIR / image_name
    if image_path.exists():
        return image_path

    # Try with URL decoding
    import urllib.parse
    decoded_name = urllib.parse.unquote(image_name)
    image_path = SCREENSHOTS_DIR / decoded_name
    if image_path.exists():
        return image_path

    # Search for partial match
    for f in SCREENSHOTS_DIR.glob("*.jpeg"):
        if image_name in f.name or decoded_name in f.name:
            return f

    for f in SCREENSHOTS_DIR.glob("*.jpg"):
        if image_name in f.name or decoded_name in f.name:
            return f

    return None

def get_figure_number():
    """Get next figure number for current section."""
    sec = current_section[0]
    if sec not in figure_counter:
        figure_counter[sec] = 0
    figure_counter[sec] += 1
    return f"{sec}.{figure_counter[sec]}"

def get_table_number():
    """Get next table number for current section."""
    sec = current_section[0]
    if sec not in table_counter:
        table_counter[sec] = 0
    table_counter[sec] += 1
    return f"{sec}.{table_counter[sec]}"

def add_figure_caption(doc, caption_text):
    """Add proper figure caption with numbering."""
    fig_num = get_figure_number()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Gambar X.Y: " in bold
    run1 = p.add_run(f"Gambar {fig_num}: ")
    run1.font.size = Pt(10)
    run1.font.bold = True
    run1.font.name = "Calibri"

    # Caption text in italic
    run2 = p.add_run(caption_text)
    run2.font.size = Pt(10)
    run2.font.italic = True
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(80, 80, 80)

    return fig_num

def add_table_caption(doc, caption_text):
    """Add proper table caption with numbering."""
    tbl_num = get_table_number()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # "Tabel X.Y: " in bold
    run1 = p.add_run(f"Tabel {tbl_num}: ")
    run1.font.size = Pt(10)
    run1.font.bold = True
    run1.font.name = "Calibri"

    # Caption text
    run2 = p.add_run(caption_text)
    run2.font.size = Pt(10)
    run2.font.name = "Calibri"
    run2.font.color.rgb = RGBColor(80, 80, 80)

    return tbl_num

def add_image(doc, image_path, caption=None):
    """Add single image with proper caption numbering."""
    if not image_path or not Path(image_path).exists():
        print(f"Warning: Image not found: {image_path}")
        return

    # Add image paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()

    try:
        run.add_picture(str(image_path), width=IMAGE_WIDTH)
    except Exception as e:
        print(f"Error adding image {image_path}: {e}")
        return

    # Add caption with proper numbering
    if caption:
        add_figure_caption(doc, caption)

    # Add small spacing after image
    doc.add_paragraph()

def add_images_side_by_side(doc, images_data):
    """Add multiple images side by side using a table (2 per row) with proper captions."""
    if not images_data:
        return

    # Calculate how many rows we need (2 images per row)
    num_images = len(images_data)
    num_rows = (num_images + 1) // 2  # Ceiling division

    # Create table with 2 columns
    table = doc.add_table(rows=num_rows * 2, cols=2)  # *2 for image + caption rows
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            cell._element.get_or_add_tcPr().append(
                OxmlElement('w:tcBorders')
            )

    # Collect figure numbers for combined caption
    fig_numbers = []
    captions = []

    img_idx = 0
    for row_idx in range(num_rows):
        img_row = table.rows[row_idx * 2]
        caption_row = table.rows[row_idx * 2 + 1]

        for col_idx in range(2):
            if img_idx < num_images:
                image_path, caption = images_data[img_idx]

                if image_path and Path(image_path).exists():
                    # Add image to cell
                    cell = img_row.cells[col_idx]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    try:
                        run.add_picture(str(image_path), width=Inches(1.8))
                    except Exception as e:
                        print(f"Error adding image {image_path}: {e}")

                    # Get figure number for this image
                    fig_num = get_figure_number()
                    fig_numbers.append(fig_num)
                    if caption:
                        captions.append(caption)

                    # Add individual caption label to cell below
                    cap_cell = caption_row.cells[col_idx]
                    cap_p = cap_cell.paragraphs[0]
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    cap_run = cap_p.add_run(f"Gambar {fig_num}")
                    cap_run.font.size = Pt(9)
                    cap_run.font.bold = True
                    cap_run.font.name = "Calibri"

                img_idx += 1

    # Add combined caption below table
    if captions:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create range like "Gambar 3.1-3.3:"
        if len(fig_numbers) > 1:
            caption_label = f"Gambar {fig_numbers[0]}-{fig_numbers[-1]}: "
        else:
            caption_label = f"Gambar {fig_numbers[0]}: "

        run1 = p.add_run(caption_label)
        run1.font.size = Pt(10)
        run1.font.bold = True
        run1.font.name = "Calibri"

        # Combined caption text
        combined_caption = captions[0] if len(captions) == 1 else " & ".join(captions[:2]) if len(captions) == 2 else captions[0]
        run2 = p.add_run(combined_caption)
        run2.font.size = Pt(10)
        run2.font.italic = True
        run2.font.name = "Calibri"
        run2.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

def add_formatted_text(paragraph, text):
    """
    Add text to paragraph with proper markdown to Word formatting conversion.
    Handles: **bold**, *italic*, ***bold italic***, __underline__, `code`
    """
    # Pattern to match markdown formatting
    # Order matters: bold italic first, then bold, then italic, then underline, then code
    patterns = [
        (r'\*\*\*([^*]+)\*\*\*', 'bold_italic'),  # ***bold italic***
        (r'\*\*([^*]+)\*\*', 'bold'),              # **bold**
        (r'\*([^*]+)\*', 'italic'),                # *italic*
        (r'__([^_]+)__', 'underline'),             # __underline__
        (r'_([^_]+)_', 'italic'),                  # _italic_
        (r'`([^`]+)`', 'code'),                    # `code`
    ]

    # Combined pattern to find all formatted segments
    combined_pattern = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*|__[^_]+__|_[^_]+_|`[^`]+`)'

    # Split text by formatted segments
    parts = re.split(combined_pattern, text)

    for part in parts:
        if not part:
            continue

        # Check which format this part matches
        format_type = None
        content = part

        # Check for bold italic (***text***)
        match = re.match(r'\*\*\*([^*]+)\*\*\*', part)
        if match:
            format_type = 'bold_italic'
            content = match.group(1)

        # Check for bold (**text**)
        if not format_type:
            match = re.match(r'\*\*([^*]+)\*\*', part)
            if match:
                format_type = 'bold'
                content = match.group(1)

        # Check for italic (*text* or _text_)
        if not format_type:
            match = re.match(r'\*([^*]+)\*', part)
            if match:
                format_type = 'italic'
                content = match.group(1)

        if not format_type:
            match = re.match(r'_([^_]+)_', part)
            if match:
                format_type = 'italic'
                content = match.group(1)

        # Check for underline (__text__)
        if not format_type:
            match = re.match(r'__([^_]+)__', part)
            if match:
                format_type = 'underline'
                content = match.group(1)

        # Check for code (`text`)
        if not format_type:
            match = re.match(r'`([^`]+)`', part)
            if match:
                format_type = 'code'
                content = match.group(1)

        # Add run with appropriate formatting
        run = paragraph.add_run(content)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

        if format_type == 'bold':
            run.font.bold = True
        elif format_type == 'italic':
            run.font.italic = True
        elif format_type == 'bold_italic':
            run.font.bold = True
            run.font.italic = True
        elif format_type == 'underline':
            run.font.underline = True
        elif format_type == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(128, 0, 0)  # Dark red for code

def add_formatted_text_to_cell(cell, text):
    """Add formatted text to a table cell."""
    p = cell.paragraphs[0]
    p.clear()
    add_formatted_text(p, text)

def parse_table(lines):
    """Parse markdown table into rows."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            # Remove outer pipes and split
            cells = [cell.strip() for cell in line[1:-1].split('|')]
            # Skip separator rows
            if not all(set(cell) <= set('-: ') for cell in cells):
                rows.append(cells)
    return rows

def add_table(doc, rows, caption=None):
    """Add table to document with proper caption numbering and formatting."""
    if not rows:
        return

    # Add table caption BEFORE table (Word standard)
    if caption:
        add_table_caption(doc, caption)
    else:
        # Generate caption from first header cell
        if rows and rows[0]:
            header_text = rows[0][0] if rows[0][0] else "Data"
            add_table_caption(doc, f"Tabel {header_text}")

    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < len(row.cells):
                cell = row.cells[j]
                p = cell.paragraphs[0]
                p.clear()

                # Convert markdown links first
                cell_text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cell_text)

                # Bold first row (header)
                if i == 0:
                    run = p.add_run(cell_text.replace('**', '').replace('*', '').replace('`', ''))
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
                else:
                    # Use formatted text for content rows
                    add_formatted_text(p, cell_text)
                    # Adjust font size for table cells
                    for run in p.runs:
                        run.font.size = Pt(10)

    doc.add_paragraph()

def process_markdown(doc, md_content):
    """Process markdown content and add to document."""
    global current_section  # Use global counter
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_content = []
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Skip HTML alignment tags
        if line.strip().startswith('<p align') or line.strip().startswith('</p>') or \
           line.strip().startswith('<h') or line.strip().startswith('</h') or \
           line.strip().startswith('<a href') or line.strip().startswith('</a>'):
            i += 1
            continue

        # Skip empty separator lines
        if line.strip() == '---':
            doc.add_paragraph()
            i += 1
            continue

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_content))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_content.append(line)
            i += 1
            continue

        # Tables
        if line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            rows = parse_table(table_lines)
            add_table(doc, rows)
            in_table = False
            table_lines = []
            # Don't increment i, process current line
            continue

        # Headings
        if line.startswith('## '):
            title = line[3:].strip()
            # Remove emojis and clean
            title = re.sub(r'[^\w\s\.\-\(\)&]', '', title).strip()

            # Extract section number from title (e.g., "1. Tentang" -> section 1)
            sec_match = re.match(r'^(\d+)\.?\s*', title)
            if sec_match:
                current_section[0] = int(sec_match.group(1))
                # Reset counters for new section
                figure_counter[current_section[0]] = 0
                table_counter[current_section[0]] = 0

            doc.add_heading(title, level=1)
            i += 1
            continue

        if line.startswith('### '):
            title = line[4:].strip()
            title = re.sub(r'[^\w\s\.\-\(\)&?]', '', title).strip()
            doc.add_heading(title, level=2)
            i += 1
            continue

        if line.startswith('#### '):
            title = line[5:].strip()
            title = re.sub(r'[^\w\s\.\-\(\)&?:]', '', title).strip()
            doc.add_heading(title, level=3)
            i += 1
            continue

        # Images - collect consecutive images and display side by side
        img_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', line)
        if img_match:
            # Collect all consecutive images
            images_to_add = []

            while i < len(lines):
                current_line = lines[i]
                current_match = re.search(r'!\[([^\]]*)\]\(([^)]+)\)', current_line)

                if current_match:
                    caption = current_match.group(1)
                    img_path = current_match.group(2)

                    # Extract filename from path
                    if '../assets/screenshots/' in img_path:
                        img_name = img_path.replace('../assets/screenshots/', '')
                    elif 'assets/screenshots/' in img_path:
                        img_name = img_path.replace('assets/screenshots/', '')
                    else:
                        img_name = os.path.basename(img_path)

                    # URL decode
                    import urllib.parse
                    img_name = urllib.parse.unquote(img_name)

                    # Find image file
                    image_file = find_image(img_name)
                    if image_file:
                        images_to_add.append((image_file, caption if caption else None))
                    else:
                        print(f"Image not found: {img_name}")

                    i += 1
                elif current_line.strip() == '':
                    # Skip empty lines between images
                    i += 1
                else:
                    # Non-image line, stop collecting
                    break

            # Add images: if multiple, show side by side; if single, show alone
            if len(images_to_add) > 1:
                add_images_side_by_side(doc, images_to_add)
            elif len(images_to_add) == 1:
                add_image(doc, images_to_add[0][0], images_to_add[0][1])

            continue

        # Bold text and lists - with proper formatting conversion
        if line.strip():
            text = line.strip()

            # Convert markdown links first (keep link text)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)

            # Handle bullet points
            if text.startswith('- ') or text.startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                add_formatted_text(p, text[2:])
            elif re.match(r'^\d+\.\s', text):
                # Numbered list
                num_text = re.sub(r'^\d+\.\s*', '', text)
                p = doc.add_paragraph(style='List Number')
                add_formatted_text(p, num_text)
            elif text.startswith('> '):
                # Quote/Note - check for callout box indicators
                quote_text = text[2:].strip()

                # Check for callout indicators
                if quote_text.startswith('💡') or quote_text.lower().startswith('tip:') or quote_text.lower().startswith('[tip]'):
                    # Tip callout
                    content = re.sub(r'^(💡|tip:|Tip:|\[tip\]|\[TIP\])\s*', '', quote_text, flags=re.IGNORECASE)
                    add_callout_box(doc, 'tip', 'Tips', content)
                elif quote_text.startswith('⚠️') or quote_text.lower().startswith('warning:') or quote_text.lower().startswith('[warning]'):
                    # Warning callout
                    content = re.sub(r'^(⚠️|warning:|Warning:|\[warning\]|\[WARNING\])\s*', '', quote_text, flags=re.IGNORECASE)
                    add_callout_box(doc, 'warning', 'Peringatan', content)
                elif quote_text.startswith('ℹ️') or quote_text.lower().startswith('note:') or quote_text.lower().startswith('[note]'):
                    # Note callout
                    content = re.sub(r'^(ℹ️|note:|Note:|\[note\]|\[NOTE\])\s*', '', quote_text, flags=re.IGNORECASE)
                    add_callout_box(doc, 'note', 'Catatan', content)
                elif quote_text.startswith('❌') or quote_text.lower().startswith('danger:') or quote_text.lower().startswith('[danger]'):
                    # Danger callout
                    content = re.sub(r'^(❌|danger:|Danger:|\[danger\]|\[DANGER\])\s*', '', quote_text, flags=re.IGNORECASE)
                    add_callout_box(doc, 'danger', 'Penting', content)
                elif quote_text.startswith('✅'):
                    # Success/tip callout
                    content = quote_text.replace('✅', '').strip()
                    add_callout_box(doc, 'tip', 'Berhasil', content)
                else:
                    # Regular quote
                    p = doc.add_paragraph()
                    add_formatted_text(p, quote_text)
                    for run in p.runs:
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(100, 100, 100)
            else:
                # Regular paragraph with formatting
                p = doc.add_paragraph()
                add_formatted_text(p, text)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        rows = parse_table(table_lines)
        add_table(doc, rows)

def add_footer_info(doc):
    """Add document footer information."""
    doc.add_page_break()

    p = doc.add_heading("Informasi Dokumen", level=1)

    p = doc.add_paragraph()
    p.add_run("Judul: ").bold = True
    p.add_run("Panduan Pengguna Gateway Config App")

    p = doc.add_paragraph()
    p.add_run("Versi: ").bold = True
    p.add_run("1.0.0")

    p = doc.add_paragraph()
    p.add_run("Tanggal: ").bold = True
    p.add_run("Desember 2025")

    p = doc.add_paragraph()
    p.add_run("Penulis: ").bold = True
    p.add_run("Tim Dokumentasi SURIOTA")

    p = doc.add_paragraph()
    p.add_run("Dibantu oleh: ").bold = True
    p.add_run("Claude AI (Anthropic)")

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Copyright 2025 PT Surya Inovasi Prioritas (SURIOTA)")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("www.suriota.com | support@suriota.com")
    run.font.size = Pt(10)

def add_sample_callouts(doc):
    """Add sample callout boxes to demonstrate the feature."""
    # This function adds callouts at strategic points in the document
    # Callouts will be added in process_markdown based on content

def main():
    """Main conversion function."""
    print("=" * 60)
    print("Converting USER_GUIDE.md to DOCX")
    print("=" * 60)

    # Create output directory
    OUTPUT_DIR.mkdir(exist_ok=True)

    # Read markdown file
    print(f"Reading: {MD_FILE}")
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    print("Creating document...")
    doc = Document()

    # Setup page margins
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN

    # Setup styles
    setup_styles(doc)

    # Setup header and footer (with page numbers)
    setup_header_footer(doc)

    # Add cover page
    print("Adding cover page...")
    add_cover_page(doc)

    # Add version history
    print("Adding version history...")
    add_version_history(doc)

    # Add table of contents
    print("Adding table of contents...")
    add_table_of_contents(doc)

    # Process markdown content
    print("Processing content...")
    process_markdown(doc, md_content)

    # Add glossary
    print("Adding glossary...")
    add_glossary(doc)

    # Add footer info
    print("Adding document info...")
    add_footer_info(doc)

    # Save document
    print(f"Saving: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)

    print("=" * 60)
    print(f"Done! Output: {OUTPUT_FILE}")
    print("=" * 60)

if __name__ == "__main__":
    main()
