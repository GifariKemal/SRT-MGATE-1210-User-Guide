#!/usr/bin/env python3
"""
enhance_docx.py - Post-process Pandoc-generated DOCX files

Enhances Pandoc-generated Word documents with SURIOTA branding:
- Professional cover page
- Page numbers in footer
- Custom headers
- Brand colors (SURIOTA teal)

Usage:
    python scripts/enhance_docx.py input.docx output.docx

Author: SURIOTA Documentation Team
Version: 1.0
Date: November 30, 2025
"""

import sys
import os
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)

# SURIOTA Brand Colors
PRIMARY_COLOR = RGBColor(0x3D, 0x7C, 0x72)  # Teal green
SECONDARY_COLOR = RGBColor(0x2C, 0x5C, 0x54)  # Dark teal
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)  # Dark gray

# Document metadata
DOC_COMPANY = "PT SURIOTA Technology Indonesia"
DOC_COPYRIGHT = "© 2025 SURIOTA. Hak Cipta Dilindungi."


def create_element(name):
    """Create an OxmlElement with the given name."""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """Set an attribute on an OxmlElement."""
    element.set(qn(name), value)


def add_page_number(paragraph):
    """Add page number field to a paragraph."""
    run = paragraph.add_run()

    # Create field character for page number
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_cover_page(doc, title="Panduan Pengguna Gateway Config App",
                   subtitle="SRT-MGATE-1210 Industrial IoT Gateway",
                   version="1.0", date=None):
    """Add SURIOTA branded cover page at the beginning."""

    if date is None:
        date = datetime.now().strftime("%d %B %Y")

    # Insert page break at beginning
    # We'll add content before first existing paragraph

    # Add spacing at top
    for _ in range(4):
        p = doc.add_paragraph()
        p.insert_paragraph_before()

    # Company logo/name
    logo_para = doc.paragraphs[0]
    logo_para.clear()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run("SURIOTA")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    # Spacing
    doc.paragraphs[1].clear()

    # Main title
    title_para = doc.paragraphs[2]
    title_para.clear()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = TEXT_COLOR

    # Subtitle
    subtitle_para = doc.paragraphs[3]
    subtitle_para.clear()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run(subtitle)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Add spacing
    for i in range(4, 10):
        if i < len(doc.paragraphs):
            doc.paragraphs[i].clear()

    # Add metadata at bottom of cover
    meta_para = doc.paragraphs[10] if len(doc.paragraphs) > 10 else doc.add_paragraph()
    meta_para.clear()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_text = f"Versi {version}\n{date}\n\n{DOC_COMPANY}"
    run = meta_para.add_run(meta_text)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    print("✓ Added cover page")


def add_headers_footers(doc):
    """Add headers and footers to all sections."""

    for section in doc.sections:
        # Header
        header = section.header
        if len(header.paragraphs) > 0:
            header_para = header.paragraphs[0]
        else:
            header_para = header.add_paragraph()

        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_para.clear()

        run = header_para.add_run("Panduan Pengguna Gateway Config App")
        run.font.size = Pt(9)
        run.font.color.rgb = PRIMARY_COLOR
        run.font.italic = True

        # Footer with page number and copyright
        footer = section.footer
        if len(footer.paragraphs) > 0:
            footer_para = footer.paragraphs[0]
        else:
            footer_para = footer.add_paragraph()

        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.clear()

        # Copyright
        run1 = footer_para.add_run(DOC_COPYRIGHT + "  |  Halaman ")
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # Page number
        add_page_number(footer_para)

    print("✓ Added headers and footers")


def apply_branding(doc):
    """Apply SURIOTA brand colors to headings."""

    # Apply colors to existing styles
    styles = doc.styles

    # Try to color Heading 1 paragraphs
    heading_count = 0
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading 1'):
            for run in para.runs:
                run.font.color.rgb = PRIMARY_COLOR
                run.font.bold = True
            heading_count += 1
        elif para.style.name.startswith('Heading 2'):
            for run in para.runs:
                run.font.color.rgb = SECONDARY_COLOR
                run.font.bold = True

    if heading_count > 0:
        print(f"✓ Applied branding to {heading_count} headings")
    else:
        print("⚠ No headings found to apply branding")


def enhance_document(input_path, output_path, add_cover=True, add_header_footer=True):
    """
    Main function to enhance document.

    Args:
        input_path: Path to input DOCX file (from Pandoc)
        output_path: Path to output enhanced DOCX file
        add_cover: Whether to add cover page (default: True)
        add_header_footer: Whether to add headers/footers (default: True)
    """

    print(f"\n{'='*60}")
    print(f"  Enhancing Document with SURIOTA Branding")
    print(f"{'='*60}\n")

    # Check input file exists
    if not os.path.exists(input_path):
        print(f"✗ Error: Input file not found: {input_path}")
        sys.exit(1)

    print(f"Loading: {input_path}")
    doc = Document(input_path)
    print(f"✓ Loaded document ({len(doc.paragraphs)} paragraphs)\n")

    # Apply enhancements
    if add_cover:
        print("Adding cover page...")
        add_cover_page(doc)
        print()

    if add_header_footer:
        print("Adding headers and footers...")
        add_headers_footers()
        print()

    print("Applying brand colors...")
    apply_branding(doc)
    print()

    # Save enhanced document
    print(f"Saving to: {output_path}")
    doc.save(output_path)

    # Get file sizes
    input_size = os.path.getsize(input_path) / 1024  # KB
    output_size = os.path.getsize(output_path) / 1024  # KB

    print(f"\n{'='*60}")
    print(f"  Enhancement Complete")
    print(f"{'='*60}")
    print(f"  Input:  {input_size:.1f} KB")
    print(f"  Output: {output_size:.1f} KB")
    print(f"{'='*60}\n")


def show_help():
    """Display help message."""
    help_text = """
Usage: python scripts/enhance_docx.py input.docx output.docx [options]

Enhance Pandoc-generated DOCX with SURIOTA branding.

Arguments:
    input.docx      Input DOCX file (from Pandoc)
    output.docx     Output enhanced DOCX file

Options:
    --no-cover      Skip adding cover page
    --no-header     Skip adding headers/footers
    --help, -h      Show this help message

Examples:
    # Full enhancement (default)
    python scripts/enhance_docx.py draft.docx final.docx

    # Skip cover page
    python scripts/enhance_docx.py draft.docx final.docx --no-cover

    # Minimal enhancement (colors only)
    python scripts/enhance_docx.py draft.docx final.docx --no-cover --no-header

Typical Workflow:
    1. Generate base document with Pandoc:
       pandoc USER_GUIDE_CLEAN.md -o draft.docx

    2. Enhance with branding:
       python scripts/enhance_docx.py draft.docx final.docx

    3. Review final.docx in Microsoft Word

Requirements:
    - python-docx: pip install python-docx

Author: SURIOTA Documentation Team
Version: 1.0
"""
    print(help_text)


def main():
    """Main entry point."""

    # Parse arguments
    args = sys.argv[1:]

    if '--help' in args or '-h' in args or len(args) < 2:
        show_help()
        sys.exit(0)

    input_path = args[0]
    output_path = args[1]

    # Parse options
    add_cover = '--no-cover' not in args
    add_header_footer = '--no-header' not in args

    # Run enhancement
    try:
        enhance_document(input_path, output_path, add_cover, add_header_footer)
        sys.exit(0)
    except Exception as e:
        print(f"\n✗ Error during enhancement: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
