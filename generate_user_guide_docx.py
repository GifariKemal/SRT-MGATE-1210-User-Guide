"""
Gateway Config App - User Guide Document Generator
Generates a professional Word document with images, tables, and styling.

Author: SURIOTA Documentation Team
Date: November 30, 2025
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Configuration
IMAGE_FOLDER = r"C:\Users\Administrator\Downloads\UI"
OUTPUT_FILE = r"C:\Users\Administrator\Music\GatewaySuriotaPOC\Documentation\Panduan_Pengguna_Gateway_Config_App.docx"

# Document metadata
DOC_TITLE = "Panduan Pengguna Gateway Config App"
DOC_SUBTITLE = "Manual Penggunaan Aplikasi Mobile untuk Konfigurasi Gateway SRT-MGATE-1210"
DOC_VERSION = "1.0"
APP_VERSION = "1.0.0"
DOC_DATE = "30 November 2025"
DOC_AUTHOR = "Tim Dokumentasi SURIOTA"
DOC_COMPANY = "PT SURIOTA Technology Indonesia"

# Color scheme (SURIOTA brand colors)
PRIMARY_COLOR = RGBColor(0x3D, 0x7C, 0x72)  # Teal green
SECONDARY_COLOR = RGBColor(0x2C, 0x5C, 0x54)  # Dark teal
TEXT_COLOR = RGBColor(0x33, 0x33, 0x33)  # Dark gray
LIGHT_COLOR = RGBColor(0xF5, 0xF5, 0xF5)  # Light gray


def create_element(name):
    """Create an OxmlElement with the given name."""
    return OxmlElement(name)


def create_attribute(element, name, value):
    """Set an attribute on an OxmlElement."""
    element.set(qn(name), value)


def add_page_number(paragraph):
    """Add page number field to a paragraph."""
    run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def setup_styles(doc):
    """Setup document styles."""
    styles = doc.styles

    # Title style
    title_style = styles['Title']
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = PRIMARY_COLOR
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)

    # Heading 1
    h1 = styles['Heading 1']
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY_COLOR
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = SECONDARY_COLOR
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = styles['Heading 3']
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = TEXT_COLOR
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Normal text
    normal = styles['Normal']
    normal.font.size = Pt(11)
    normal.font.name = 'Calibri'
    normal.font.color.rgb = TEXT_COLOR
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(8)

    return styles


def add_header_footer(doc):
    """Add header and footer to all sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        run = header_para.add_run(DOC_TITLE)
        run.font.size = Pt(9)
        run.font.color.rgb = PRIMARY_COLOR
        run.font.italic = True

        # Footer with page number
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run1 = footer_para.add_run("© 2025 SURIOTA  |  ")
        run1.font.size = Pt(9)
        run1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        run2 = footer_para.add_run("Halaman ")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        add_page_number(footer_para)


def add_cover_page(doc):
    """Create cover page."""
    # Add spacing at top
    for _ in range(4):
        doc.add_paragraph()

    # Company logo placeholder
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = logo_para.add_run("SURIOTA")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(DOC_TITLE)
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = TEXT_COLOR

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(DOC_SUBTITLE)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Spacing
    for _ in range(6):
        doc.add_paragraph()

    # Document info box
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("Versi Dokumen", DOC_VERSION),
        ("Versi Aplikasi", APP_VERSION),
        ("Tanggal Terbit", DOC_DATE),
        ("Penulis", DOC_AUTHOR),
        ("Penerbit", DOC_COMPANY)
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(3)

    # Page break
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add table of contents."""
    doc.add_heading("Daftar Isi", level=1)

    toc_items = [
        ("1.", "Pendahuluan", "3"),
        ("2.", "Persyaratan Sistem", "4"),
        ("3.", "Instalasi Aplikasi", "5"),
        ("4.", "Memulai Aplikasi", "6"),
        ("5.", "Menghubungkan Gateway via Bluetooth", "7"),
        ("6.", "Menu Utama Gateway", "10"),
        ("7.", "Mengatur Perangkat Komunikasi", "12"),
        ("8.", "Mengatur Data Register", "16"),
        ("9.", "Mengatur Koneksi Server", "19"),
        ("10.", "Status Perangkat", "24"),
        ("11.", "Pengaturan & Backup", "25"),
        ("12.", "Pengaturan Aplikasi", "27"),
        ("13.", "Pemecahan Masalah", "28"),
        ("14.", "Tanya Jawab (FAQ)", "30"),
        ("15.", "Lampiran", "32"),
        ("16.", "Kontak & Dukungan", "34"),
    ]

    toc_table = doc.add_table(rows=len(toc_items), cols=3)

    for i, (num, title, page) in enumerate(toc_items):
        row = toc_table.rows[i]
        row.cells[0].text = num
        row.cells[1].text = title
        row.cells[2].text = page
        row.cells[0].width = Inches(0.5)
        row.cells[1].width = Inches(4.5)
        row.cells[2].width = Inches(0.5)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()


def add_styled_table(doc, headers, data, header_color="3D7C72"):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, header_color)

    # Data rows
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        for col_idx, cell_data in enumerate(row_data):
            row.cells[col_idx].text = str(cell_data)

    doc.add_paragraph()
    return table


def add_image_with_caption(doc, image_path, caption, width=Inches(3)):
    """Add an image with a caption below it."""
    if os.path.exists(image_path):
        # Center the image
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(image_path, width=width)

        # Add caption
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()
        return True
    return False


def add_note_box(doc, text, note_type="info"):
    """Add a highlighted note/tip box."""
    colors = {
        "info": ("E8F4F8", "2196F3"),
        "warning": ("FFF3E0", "FF9800"),
        "tip": ("E8F5E9", "4CAF50"),
    }
    bg_color, border_color = colors.get(note_type, colors["info"])

    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    cell.text = text
    set_cell_shading(cell, bg_color)

    doc.add_paragraph()


def get_sorted_images():
    """Get all images sorted by timestamp in filename."""
    images = []
    for f in os.listdir(IMAGE_FOLDER):
        if f.lower().endswith(('.jpeg', '.jpg', '.png')):
            images.append(os.path.join(IMAGE_FOLDER, f))
    return sorted(images)


def add_section_1(doc, images):
    """Section 1: Pendahuluan"""
    doc.add_heading("1. Pendahuluan", level=1)

    doc.add_heading("1.1 Tentang Panduan Ini", level=2)
    doc.add_paragraph(
        "Panduan ini ditujukan untuk membantu pengguna dalam mengoperasikan aplikasi "
        "Gateway Config App untuk mengkonfigurasi perangkat gateway SRT-MGATE-1210. "
        "Panduan ini ditulis untuk pengguna umum tanpa memerlukan pengetahuan teknis mendalam."
    )

    doc.add_heading("1.2 Tentang Aplikasi", level=2)
    doc.add_paragraph("Gateway Config App adalah aplikasi smartphone yang memungkinkan Anda untuk:")

    features = [
        "Menghubungkan smartphone ke gateway melalui Bluetooth",
        "Mengatur koneksi jaringan (WiFi atau Kabel LAN)",
        "Menambahkan dan mengatur perangkat sensor/meter",
        "Mengkonfigurasi pengiriman data ke server cloud",
        "Melakukan backup dan restore pengaturan",
        "Memperbarui firmware gateway"
    ]
    for feature in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(feature)

    doc.add_heading("1.3 Tentang Gateway SRT-MGATE-1210", level=2)
    doc.add_paragraph(
        "Gateway SRT-MGATE-1210 adalah perangkat penghubung yang mengumpulkan data dari "
        "berbagai sensor dan meter industri, kemudian mengirimkannya ke server cloud "
        "untuk dipantau dan dianalisis."
    )

    doc.add_paragraph("Kemampuan Gateway:")
    capabilities = [
        "Membaca data dari perangkat industri (sensor suhu, meter listrik, dll)",
        "Mengirim data melalui WiFi atau kabel Ethernet",
        "Mendukung protokol MQTT dan HTTP untuk pengiriman data",
        "Dapat dikonfigurasi tanpa kabel melalui Bluetooth"
    ]
    for cap in capabilities:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(cap)

    doc.add_page_break()


def add_section_2(doc, images):
    """Section 2: Persyaratan Sistem"""
    doc.add_heading("2. Persyaratan Sistem", level=1)

    doc.add_heading("2.1 Persyaratan Smartphone", level=2)

    doc.add_paragraph("Untuk Pengguna Android:", style='Heading 3')
    add_styled_table(doc,
        ["Persyaratan", "Spesifikasi"],
        [
            ["Sistem Operasi", "Android versi 5.0 atau lebih baru"],
            ["Bluetooth", "BLE 4.0 atau lebih tinggi"],
            ["Penyimpanan", "Minimal 100 MB ruang kosong"],
            ["Koneksi Internet", "Untuk download aplikasi"]
        ]
    )

    doc.add_paragraph("Untuk Pengguna iPhone/iPad:", style='Heading 3')
    add_styled_table(doc,
        ["Persyaratan", "Spesifikasi"],
        [
            ["Sistem Operasi", "iOS versi 12 atau lebih baru"],
            ["Bluetooth", "Aktif dan didukung"],
            ["Penyimpanan", "Minimal 100 MB ruang kosong"]
        ]
    )

    doc.add_heading("2.2 Persyaratan Gateway", level=2)
    add_styled_table(doc,
        ["Item", "Keterangan"],
        [
            ["Model", "SRT-MGATE-1210"],
            ["Firmware", "Versi 2.0.0 atau lebih baru"],
            ["Status", "Menyala dan siap digunakan"]
        ]
    )

    doc.add_heading("2.3 Izin yang Diperlukan", level=2)
    add_styled_table(doc,
        ["Izin", "Kegunaan"],
        [
            ["Bluetooth", "Untuk menghubungkan ke gateway"],
            ["Lokasi", "Diperlukan untuk pencarian perangkat Bluetooth"],
            ["Penyimpanan", "Untuk menyimpan file backup"]
        ]
    )

    doc.add_page_break()


def add_section_3(doc, images):
    """Section 3: Instalasi Aplikasi"""
    doc.add_heading("3. Instalasi Aplikasi", level=1)

    doc.add_heading("3.1 Untuk Pengguna Android", level=2)
    steps_android = [
        "Buka aplikasi Play Store di smartphone Anda",
        'Ketik "Gateway Config" atau "SURIOTA Gateway" di kolom pencarian',
        "Pilih aplikasi yang sesuai dari hasil pencarian",
        "Tekan tombol Install",
        "Tunggu hingga proses instalasi selesai",
        "Tekan Open untuk membuka aplikasi"
    ]
    for i, step in enumerate(steps_android, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)

    doc.add_heading("3.2 Untuk Pengguna iPhone/iPad", level=2)
    steps_ios = [
        "Buka aplikasi App Store di perangkat Anda",
        'Ketik "Gateway Config" atau "SURIOTA Gateway" di kolom pencarian',
        "Pilih aplikasi yang sesuai dari hasil pencarian",
        "Tekan tombol Get (atau ikon download)",
        "Verifikasi dengan Face ID, Touch ID, atau password Apple ID",
        "Tunggu hingga proses instalasi selesai"
    ]
    for i, step in enumerate(steps_ios, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)

    doc.add_heading("3.3 Pemberian Izin Pertama Kali", level=2)
    doc.add_paragraph("Saat pertama kali membuka aplikasi:")
    permissions = [
        ("Bluetooth", "Tekan Izinkan"),
        ("Lokasi", "Tekan Izinkan"),
        ("Penyimpanan", "Tekan Izinkan")
    ]
    for perm, action in permissions:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"Permintaan izin {perm} - {action}")

    add_note_box(doc, "Catatan: Jika Anda tidak sengaja menolak izin, Anda dapat mengaktifkannya "
                     "kembali melalui menu Pengaturan smartphone.", "info")

    doc.add_page_break()


def add_section_4(doc, images):
    """Section 4: Memulai Aplikasi"""
    doc.add_heading("4. Memulai Aplikasi", level=1)

    doc.add_heading("4.1 Halaman Utama (Home)", level=2)
    doc.add_paragraph("Setelah membuka aplikasi, Anda akan melihat halaman utama yang menampilkan:")

    items = [
        ("Daftar Perangkat", "Gateway yang pernah dihubungkan"),
        ("Tombol Tambah (+)", "Untuk mencari dan menghubungkan gateway baru"),
        ("Menu Settings", "Pengaturan aplikasi")
    ]
    for item, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        p.add_run(f" - {desc}")

    # Add Home screen image
    home_images = [img for img in images if "06.13.24 (1)" in img or "06.13.24 (2)" in img]
    if home_images:
        add_image_with_caption(doc, home_images[0], "Gambar 4.1 - Halaman Utama Aplikasi")

    doc.add_heading("4.2 Navigasi Aplikasi", level=2)
    doc.add_paragraph("Aplikasi memiliki dua tab utama di bagian bawah layar:")
    add_styled_table(doc,
        ["Tab", "Fungsi"],
        [
            ["Home", "Melihat dan mengelola gateway yang terhubung"],
            ["Settings", "Mengatur profil dan preferensi aplikasi"]
        ]
    )

    doc.add_page_break()


def add_section_5(doc, images):
    """Section 5: Menghubungkan Gateway via Bluetooth"""
    doc.add_heading("5. Menghubungkan Gateway via Bluetooth", level=1)

    doc.add_heading("5.1 Persiapan Gateway", level=2)
    doc.add_paragraph("Sebelum menghubungkan, pastikan:")
    preps = [
        "Gateway dalam keadaan menyala - Periksa lampu indikator",
        "Bluetooth gateway aktif - Tekan tombol pada gateway untuk mengaktifkan Bluetooth",
        "Bluetooth akan aktif selama 5 menit setelah tombol ditekan"
    ]
    for i, prep in enumerate(preps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(prep)

    doc.add_heading("5.2 Langkah-langkah Menghubungkan", level=2)

    # Step 1
    doc.add_paragraph("Langkah 1: Mulai Pencarian", style='Heading 3')
    step1 = [
        "Di halaman utama aplikasi, tekan tombol [+]",
        "Anda akan masuk ke halaman Scan Devices",
        "Tekan tombol Scan untuk memulai pencarian"
    ]
    for s in step1:
        doc.add_paragraph(s, style='List Bullet')

    # Add scan device image
    scan_images = [img for img in images if "06.13.24.jpeg" in img]
    if scan_images:
        add_image_with_caption(doc, scan_images[0], "Gambar 5.1 - Halaman Scan Devices")

    # Step 2
    doc.add_paragraph("Langkah 2: Proses Pencarian", style='Heading 3')
    doc.add_paragraph("Aplikasi akan menampilkan animasi loading dengan pesan \"Scanning device...\"")

    scanning_images = [img for img in images if "06.13.23 (2)" in img]
    if scanning_images:
        add_image_with_caption(doc, scanning_images[0], "Gambar 5.2 - Proses Scanning")

    add_note_box(doc, "Tips: Proses pencarian biasanya memakan waktu 5-10 detik. "
                     "Pastikan Anda berada dalam jarak 10 meter dari gateway.", "tip")

    # Step 3
    doc.add_paragraph("Langkah 3: Pilih Gateway", style='Heading 3')
    doc.add_paragraph("Setelah pencarian selesai, daftar perangkat Bluetooth akan muncul:")
    step3 = [
        'Cari gateway dengan nama "SURIOTA GW"',
        "Gunakan kolom pencarian untuk memfilter hasil",
        "Tekan tombol Connect pada gateway yang diinginkan"
    ]
    for s in step3:
        doc.add_paragraph(s, style='List Bullet')

    results_images = [img for img in images if "06.13.23 (1)" in img]
    if results_images:
        add_image_with_caption(doc, results_images[0], "Gambar 5.3 - Hasil Pencarian Perangkat")

    # Step 4
    doc.add_paragraph("Langkah 4: Konfirmasi Koneksi", style='Heading 3')
    doc.add_paragraph("Setelah berhasil terhubung, dialog konfirmasi akan muncul. "
                     "Tekan Yes untuk masuk ke halaman pengaturan gateway.")

    connected_images = [img for img in images if "06.13.22 (2)" in img]
    if connected_images:
        add_image_with_caption(doc, connected_images[0], "Gambar 5.4 - Dialog Konfirmasi Koneksi")

    doc.add_heading("5.3 Status Koneksi", level=2)
    doc.add_paragraph("Setelah terhubung:")
    status_items = [
        ("BONDED", "Gateway terpasang dan terhubung (ditandai warna hijau)"),
        ("Disconnect", "Tombol untuk memutuskan koneksi (warna merah)")
    ]
    for status, desc in status_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(status)
        run.bold = True
        p.add_run(f" - {desc}")

    doc.add_page_break()


def add_section_6(doc, images):
    """Section 6: Menu Utama Gateway"""
    doc.add_heading("6. Menu Utama Gateway", level=1)

    doc.add_heading("6.1 Dashboard Gateway", level=2)
    doc.add_paragraph("Setelah terhubung, Anda akan melihat dashboard dengan informasi:")
    info_items = [
        ("Nama Gateway", "Contoh: SURIOTA GW"),
        ("Status Koneksi", "BONDED (terhubung)"),
        ("ID Perangkat", "Alamat unik gateway"),
        ("Tombol Disconnect", "Untuk memutuskan koneksi")
    ]
    for item, desc in info_items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.bold = True
        p.add_run(f" - {desc}")

    detail_images = [img for img in images if "06.13.22 (1)" in img]
    if detail_images:
        add_image_with_caption(doc, detail_images[0], "Gambar 6.1 - Dashboard Detail Device")

    doc.add_heading("6.2 Menu Konfigurasi", level=2)
    doc.add_paragraph("Dashboard menampilkan 6 menu utama:")
    add_styled_table(doc,
        ["Menu", "Fungsi"],
        [
            ["Device Communications", "Mengatur sensor/meter yang terhubung"],
            ["Modbus Configurations", "Mengatur data yang dibaca dari sensor"],
            ["Server Configurations", "Mengatur koneksi jaringan dan server"],
            ["Logging Configurations", "Mengatur penyimpanan log (Segera Hadir)"],
            ["Status", "Melihat status dan update firmware"],
            ["Settings", "Backup, restore, dan reset"]
        ]
    )

    doc.add_page_break()


def add_section_7(doc, images):
    """Section 7: Mengatur Perangkat Komunikasi"""
    doc.add_heading("7. Mengatur Perangkat Komunikasi", level=1)

    doc.add_paragraph("Menu Device Communications digunakan untuk menambah dan mengatur "
                     "perangkat sensor atau meter yang terhubung ke gateway.")

    doc.add_heading("7.1 Melihat Daftar Perangkat", level=2)
    doc.add_paragraph("Halaman ini menampilkan:")
    items = [
        "Daftar semua perangkat yang sudah dikonfigurasi",
        "Status aktif/nonaktif setiap perangkat",
        "Jumlah data (register) yang dikonfigurasi",
        "Tombol untuk melihat, mengedit, dan menghapus"
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    device_list_images = [img for img in images if "06.13.18.jpeg" in img]
    if device_list_images:
        add_image_with_caption(doc, device_list_images[0], "Gambar 7.1 - Daftar Perangkat")

    doc.add_heading("7.2 Menambah Perangkat Baru", level=2)
    doc.add_paragraph("Untuk menambah perangkat baru:")
    steps = [
        "Tekan tombol [+] di pojok kanan atas",
        "Isi informasi perangkat pada form yang muncul"
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)

    doc.add_paragraph("Informasi Dasar:", style='Heading 3')
    add_styled_table(doc,
        ["Field", "Keterangan", "Contoh"],
        [
            ["Device Name", "Nama untuk identifikasi", "Sensor_Suhu_Ruang1"],
            ["Slave ID", "Alamat perangkat (1-247)", "1"]
        ]
    )

    doc.add_paragraph("Pilih Jenis Koneksi:", style='Heading 3')
    add_styled_table(doc,
        ["Jenis", "Keterangan"],
        [
            ["Modbus RTU", "Untuk perangkat via kabel serial (RS485)"],
            ["Modbus TCP", "Untuk perangkat via kabel jaringan (Ethernet)"]
        ]
    )

    form_images = [img for img in images if "06.13.17.jpeg" in img]
    if form_images:
        add_image_with_caption(doc, form_images[0], "Gambar 7.2 - Form Setup Device")

    doc.add_heading("7.3 Pengaturan Modbus RTU", level=2)
    doc.add_paragraph("Jika memilih Modbus RTU, atur parameter berikut:")
    add_styled_table(doc,
        ["Parameter", "Keterangan", "Nilai Umum"],
        [
            ["Serial Port", "Port yang digunakan", "PORT1 atau PORT2"],
            ["Baudrate", "Kecepatan komunikasi", "9600, 19200, 115200"],
            ["Bit Data", "Jumlah bit data", "8"],
            ["Parity", "Paritas", "None, Even, Odd"],
            ["Stop Bit", "Stop bit", "1 atau 2"]
        ]
    )

    doc.add_paragraph("Pengaturan Lanjutan:", style='Heading 3')
    add_styled_table(doc,
        ["Parameter", "Keterangan", "Rekomendasi"],
        [
            ["Retry Count", "Jumlah percobaan ulang", "3"],
            ["Timeout", "Batas waktu tunggu (milidetik)", "3000"],
            ["Refresh Rate", "Interval pembacaan (milidetik)", "5000"]
        ]
    )

    doc.add_heading("7.4 Pengaturan Modbus TCP", level=2)
    doc.add_paragraph("Jika memilih Modbus TCP, atur:")
    add_styled_table(doc,
        ["Parameter", "Keterangan", "Contoh"],
        [
            ["IP Address", "Alamat IP perangkat", "192.168.1.100"],
            ["Server Port", "Port komunikasi", "502"]
        ]
    )

    doc.add_page_break()


def add_section_8(doc, images):
    """Section 8: Mengatur Data Register"""
    doc.add_heading("8. Mengatur Data Register", level=1)

    doc.add_paragraph("Menu Modbus Configurations digunakan untuk mengatur data apa saja "
                     "yang akan dibaca dari setiap perangkat.")

    doc.add_heading("8.1 Memilih Perangkat", level=2)
    steps = [
        "Buka menu Modbus Configurations",
        "Pilih perangkat dari dropdown Choose Device",
        "Daftar register (data) yang sudah dikonfigurasi akan muncul"
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)

    modbus_images = [img for img in images if "06.13.16.jpeg" in img]
    if modbus_images:
        add_image_with_caption(doc, modbus_images[0], "Gambar 8.1 - Pilih Perangkat")

    doc.add_heading("8.2 Daftar Register", level=2)
    doc.add_paragraph("Setiap kartu register menampilkan:")
    items = ["Nama data", "Tipe data", "Alamat register", "Tombol edit dan hapus"]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("8.3 Menambah Register Baru", level=2)
    doc.add_paragraph("Tekan tombol [+] dan isi informasi register:")

    add_styled_table(doc,
        ["Field", "Keterangan", "Contoh"],
        [
            ["Data Name", "Nama data", "Suhu_Ruangan"],
            ["Choose Device", "Perangkat sumber", "Sensor_Suhu_Ruang1"],
            ["Description", "Deskripsi (opsional)", "Suhu ruang produksi"],
            ["Unit", "Satuan", "°C, %, kWh"],
            ["Choose Function", "Jenis pembacaan", "Input Registers"],
            ["Address Modbus", "Alamat data", "0"],
            ["Choose Data Type", "Format data", "INT16, FLOAT32"]
        ]
    )

    setup_modbus_images = [img for img in images if "06.13.15.jpeg" in img]
    if setup_modbus_images:
        add_image_with_caption(doc, setup_modbus_images[0], "Gambar 8.2 - Setup Modbus Register")

    doc.add_heading("8.4 Jenis Pembacaan (Function)", level=2)
    add_styled_table(doc,
        ["Jenis", "Kode", "Kegunaan"],
        [
            ["Coil Status", "01", "Membaca status ON/OFF"],
            ["Input Status", "02", "Membaca input digital"],
            ["Holding Register", "03", "Membaca/menulis data"],
            ["Input Registers", "04", "Membaca data sensor"]
        ]
    )

    doc.add_heading("8.5 Format Data (Data Type)", level=2)
    add_styled_table(doc,
        ["Format", "Keterangan"],
        [
            ["INT16", "Bilangan bulat (-32768 s/d 32767)"],
            ["UINT16", "Bilangan bulat positif (0 s/d 65535)"],
            ["FLOAT32", "Bilangan desimal"],
            ["INT32", "Bilangan bulat besar"]
        ]
    )

    add_note_box(doc, "Catatan tentang Format:\n"
                     "• BE (Big Endian) - Format standar, coba ini terlebih dahulu\n"
                     "• LE (Little Endian) - Jika data terbaca salah, coba format ini\n"
                     "• BS (Byte Swap) - Untuk perangkat dengan urutan byte berbeda", "info")

    doc.add_heading("8.6 Kalibrasi Data", level=2)
    doc.add_paragraph("Jika data perlu dikonversi, gunakan pengaturan kalibrasi:")
    add_styled_table(doc,
        ["Parameter", "Keterangan", "Contoh"],
        [
            ["Scale", "Pengali", "0.1"],
            ["Offset", "Penambah", "0"]
        ]
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Rumus: ")
    run.bold = True
    p.add_run("Nilai Akhir = (Nilai Mentah × Scale) + Offset")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Contoh: ")
    run.bold = True
    p.add_run("Sensor memberikan nilai 250, Scale = 0.1, Offset = 0")
    doc.add_paragraph("Hasil = 250 × 0.1 + 0 = 25.0°C")

    doc.add_page_break()


def add_section_9(doc, images):
    """Section 9: Mengatur Koneksi Server"""
    doc.add_heading("9. Mengatur Koneksi Server", level=1)

    doc.add_paragraph("Menu Server Configurations digunakan untuk mengatur bagaimana "
                     "gateway mengirim data ke server cloud.")

    doc.add_heading("9.1 Pengaturan Jaringan", level=2)
    doc.add_paragraph("Pilih Mode Koneksi:", style='Heading 3')
    add_styled_table(doc,
        ["Mode", "Keterangan"],
        [
            ["WIFI", "Menggunakan jaringan WiFi"],
            ["ETH", "Menggunakan kabel Ethernet"]
        ]
    )

    server_images = [img for img in images if "06.13.21.jpeg" in img]
    if server_images:
        add_image_with_caption(doc, server_images[0], "Gambar 9.1 - Pengaturan Server")

    doc.add_paragraph("Pengaturan WiFi:", style='Heading 3')
    add_styled_table(doc,
        ["Field", "Keterangan"],
        [
            ["WiFi Enabled", "Aktifkan WiFi"],
            ["WiFi SSID", "Nama jaringan WiFi"],
            ["WiFi Password", "Kata sandi WiFi"]
        ]
    )

    doc.add_paragraph("Pengaturan Ethernet:", style='Heading 3')
    add_styled_table(doc,
        ["Field", "Keterangan"],
        [
            ["Ethernet Enabled", "Aktifkan Ethernet"],
            ["Using DHCP", "Otomatis mendapat IP"],
            ["Static IP", "Alamat IP manual (jika DHCP nonaktif)"],
            ["Gateway", "Alamat gateway jaringan"],
            ["Subnet", "Subnet mask"]
        ]
    )

    doc.add_heading("9.2 Pengaturan MQTT", level=2)
    doc.add_paragraph("MQTT adalah protokol untuk mengirim data ke server cloud.")
    add_styled_table(doc,
        ["Field", "Keterangan", "Contoh"],
        [
            ["Enabled", "Aktifkan MQTT", "true"],
            ["Broker Address", "Alamat server MQTT", "broker.hivemq.com"],
            ["Broker Port", "Port server", "1883"],
            ["Client ID", "ID unik gateway", "SRT-MGATE-001"],
            ["Username", "Nama pengguna", "(opsional)"],
            ["Password", "Kata sandi", "(opsional)"],
            ["Keep Alive", "Interval ping (detik)", "60"],
            ["Use TLS", "Enkripsi data", "false"]
        ]
    )

    doc.add_heading("9.3 Pengaturan Interval Pengiriman", level=2)
    doc.add_paragraph("Interval menentukan seberapa sering data dikirim ke server.")

    doc.add_paragraph("Satuan Waktu:", style='Heading 3')
    add_styled_table(doc,
        ["Satuan", "Singkatan", "Konversi"],
        [
            ["Milidetik", "ms", "1000 ms = 1 detik"],
            ["Detik", "s", "60 s = 1 menit"],
            ["Menit", "m", "1 m = 60 detik"]
        ]
    )

    doc.add_paragraph("Tabel Konversi:", style='Heading 3')
    add_styled_table(doc,
        ["Nilai", "Satuan", "Sama Dengan"],
        [
            ["5", "s", "5 detik"],
            ["60", "s", "1 menit"],
            ["5", "m", "5 menit"],
            ["1000", "ms", "1 detik"],
            ["5000", "ms", "5 detik"]
        ]
    )

    doc.add_paragraph("Rekomendasi Interval:", style='Heading 3')
    add_styled_table(doc,
        ["Kegunaan", "Interval Rekomendasi"],
        [
            ["Monitoring real-time", "1-5 detik"],
            ["Logging standar", "30 detik - 1 menit"],
            ["Hemat bandwidth", "5-15 menit"]
        ]
    )

    doc.add_heading("9.4 Pengaturan Topic", level=2)
    doc.add_paragraph("Mode Default (Sederhana):", style='Heading 3')
    doc.add_paragraph("Semua data dikirim ke satu alamat topic. Cocok untuk penggunaan umum.")

    doc.add_paragraph("Mode Customize (Lanjutan):", style='Heading 3')
    features = [
        "Bisa membuat beberapa topic berbeda",
        "Setiap topic bisa memiliki interval berbeda",
        "Bisa memilih data spesifik untuk setiap topic"
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    custom_topic_images = [img for img in images if "06.13.20.jpeg" in img]
    if custom_topic_images:
        add_image_with_caption(doc, custom_topic_images[0], "Gambar 9.2 - Customize Topic")

    doc.add_page_break()


def add_section_10(doc, images):
    """Section 10: Status Perangkat"""
    doc.add_heading("10. Status Perangkat", level=1)

    doc.add_paragraph("Menu Status menampilkan informasi tentang gateway.")

    doc.add_heading("10.1 Informasi yang Ditampilkan", level=2)
    add_styled_table(doc,
        ["Item", "Keterangan"],
        [
            ["Firmware", "Versi perangkat lunak gateway"],
            ["SD Card Info", "Informasi kartu memori (Segera Hadir)"],
            ["Update Firmware", "Memperbarui perangkat lunak gateway"]
        ]
    )

    doc.add_heading("10.2 Memperbarui Firmware", level=2)
    steps = [
        "Tekan menu Update Firmware",
        "Pilih file firmware (.bin) dari penyimpanan",
        "Tunggu proses update selesai (jangan matikan gateway!)",
        "Gateway akan restart otomatis"
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)

    add_note_box(doc, "PERINGATAN: Jangan matikan gateway atau menutup aplikasi "
                     "selama proses update berlangsung.", "warning")

    doc.add_page_break()


def add_section_11(doc, images):
    """Section 11: Pengaturan & Backup"""
    doc.add_heading("11. Pengaturan & Backup", level=1)

    doc.add_paragraph("Menu Settings pada dashboard gateway menyediakan fitur untuk mengelola konfigurasi.")

    settings_images = [img for img in images if "06.13.22.jpeg" in img and "06.13.22 (" not in img]
    if settings_images:
        add_image_with_caption(doc, settings_images[0], "Gambar 11.1 - Device Settings")

    doc.add_heading("11.1 Import Config (Memuat Konfigurasi)", level=2)
    doc.add_paragraph("Untuk memuat konfigurasi dari file backup:")
    steps = [
        "Tekan Import Config",
        "Pilih file .json dari penyimpanan smartphone",
        "Konfirmasi untuk memuat konfigurasi",
        "Gateway akan menerapkan pengaturan dari file"
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)

    doc.add_heading("11.2 Download All Config (Backup)", level=2)
    doc.add_paragraph("Untuk menyimpan semua konfigurasi ke file:")
    steps = [
        "Tekan Download All Config",
        "File akan tersimpan di folder: Documents/GatewayConfig/backup/",
        "Nama file: config_backup_[tanggal].json"
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Data yang di-backup:")
    run.bold = True

    backup_items = [
        "Pengaturan perangkat",
        "Pengaturan register",
        "Pengaturan jaringan",
        "Pengaturan MQTT/HTTP"
    ]
    for item in backup_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading("11.3 Clear Configuration (Reset)", level=2)
    doc.add_paragraph("Untuk menghapus semua konfigurasi dan mengembalikan ke pengaturan awal:")
    steps = [
        "Tekan Clear Configuration",
        "Konfirmasi penghapusan",
        "Semua data akan dihapus",
        "Gateway akan restart"
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. ").bold = True
        p.add_run(step)

    add_note_box(doc, "PERINGATAN: Pastikan sudah melakukan backup sebelum reset! "
                     "Data yang dihapus tidak dapat dikembalikan.", "warning")

    doc.add_page_break()


def add_section_12(doc, images):
    """Section 12: Pengaturan Aplikasi"""
    doc.add_heading("12. Pengaturan Aplikasi", level=1)

    doc.add_paragraph("Tab Settings di bagian bawah aplikasi menyediakan pengaturan umum.")

    doc.add_heading("12.1 Menu yang Tersedia", level=2)

    doc.add_paragraph("Akun:", style='Heading 3')
    doc.add_paragraph("Profile - Mengelola informasi akun pengguna", style='List Bullet')

    doc.add_paragraph("Tentang:", style='Heading 3')
    about_items = [
        "About Product - Informasi tentang gateway",
        "About App - Informasi tentang aplikasi"
    ]
    for item in about_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph("Lainnya:", style='Heading 3')
    other_items = [
        "Contact Us - Menghubungi tim dukungan",
        "Logout - Keluar dari akun"
    ]
    for item in other_items:
        doc.add_paragraph(item, style='List Bullet')

    settings_app_images = [img for img in images if "06.13.23.jpeg" in img and "06.13.23 (" not in img]
    if settings_app_images:
        add_image_with_caption(doc, settings_app_images[0], "Gambar 12.1 - Pengaturan Aplikasi")

    doc.add_page_break()


def add_section_13(doc, images):
    """Section 13: Pemecahan Masalah"""
    doc.add_heading("13. Pemecahan Masalah", level=1)

    doc.add_heading("13.1 Masalah Koneksi Bluetooth", level=2)
    add_styled_table(doc,
        ["Masalah", "Solusi"],
        [
            ["Gateway tidak ditemukan", "1. Pastikan Bluetooth gateway aktif (tekan tombol)\n"
                                        "2. Dekatkan smartphone (maks 10 meter)\n"
                                        "3. Restart Bluetooth smartphone\n"
                                        "4. Cek izin aplikasi"],
            ["Koneksi terputus", "1. Periksa jarak ke gateway\n"
                                 "2. Hindari gangguan WiFi\n"
                                 "3. Restart aplikasi"],
            ["Gagal pairing", "1. Hapus gateway dari daftar Bluetooth\n"
                             "2. Restart gateway\n"
                             "3. Coba hubungkan ulang"]
        ]
    )

    doc.add_heading("13.2 Masalah Konfigurasi", level=2)
    add_styled_table(doc,
        ["Masalah", "Solusi"],
        [
            ["Pengaturan tidak tersimpan", "1. Pastikan koneksi stabil\n"
                                           "2. Tunggu konfirmasi \"Saved\"\n"
                                           "3. Dekatkan smartphone ke gateway"],
            ["Data sensor tidak terbaca", "1. Periksa alamat register\n"
                                          "2. Verifikasi jenis pembacaan\n"
                                          "3. Cek koneksi fisik sensor"]
        ]
    )

    doc.add_heading("13.3 Masalah Jaringan", level=2)
    add_styled_table(doc,
        ["Masalah", "Solusi"],
        [
            ["WiFi tidak terhubung", "1. Periksa nama dan password WiFi\n"
                                     "2. Pastikan router menyala\n"
                                     "3. Cek jarak ke router"],
            ["Ethernet tidak aktif", "1. Periksa kabel LAN\n"
                                     "2. Verifikasi pengaturan IP\n"
                                     "3. Cek DHCP server"],
            ["Data tidak terkirim ke server", "1. Periksa alamat broker/server\n"
                                              "2. Verifikasi username/password\n"
                                              "3. Cek koneksi internet"]
        ]
    )

    doc.add_page_break()


def add_section_14(doc, images):
    """Section 14: Tanya Jawab (FAQ)"""
    doc.add_heading("14. Tanya Jawab (FAQ)", level=1)

    faqs = [
        ("Koneksi & Bluetooth", [
            ("Berapa jarak maksimal koneksi Bluetooth?",
             "Sekitar 10 meter dalam ruangan tanpa penghalang."),
            ("Gateway tidak muncul saat scan, apa yang harus dilakukan?",
             "Tekan tombol pada gateway untuk mengaktifkan Bluetooth, lalu scan ulang."),
            ("Apakah bisa mengatur gateway tanpa internet?",
             "Ya, konfigurasi via Bluetooth tidak memerlukan internet.")
        ]),
        ("Data & Pengiriman", [
            ("Apa itu interval pengiriman?",
             "Waktu jeda antar pengiriman data ke server. Misalnya 5 detik artinya data dikirim setiap 5 detik."),
            ("Apa arti ms, s, dan m pada interval?",
             "ms = milidetik (1000 ms = 1 detik)\n"
             "s = detik (60 s = 1 menit)\n"
             "m = menit (1 m = 60 detik)"),
            ("Berapa interval yang direkomendasikan?",
             "Untuk monitoring: 5 detik. Untuk logging: 1-5 menit. Untuk hemat bandwidth: 10-15 menit."),
            ("Apa yang terjadi jika internet putus?",
             "Data akan disimpan sementara dan dikirim ulang setelah internet pulih.")
        ]),
        ("Format Data", [
            ("Apa perbedaan BE dan LE?",
             "BE (Big Endian) dan LE (Little Endian) adalah urutan penyimpanan data. "
             "Coba BE terlebih dahulu, jika data salah gunakan LE."),
            ("Kapan menggunakan format dengan BS (Byte Swap)?",
             "Jika data masih salah setelah mencoba BE dan LE, coba format dengan BS.")
        ]),
        ("Backup & Reset", [
            ("Dimana file backup tersimpan?",
             "Di folder Documents/GatewayConfig/backup/ pada smartphone."),
            ("Apakah bisa memindahkan konfigurasi ke gateway lain?",
             "Ya, gunakan fitur Download Config dari gateway lama, lalu Import Config ke gateway baru.")
        ])
    ]

    for category, questions in faqs:
        doc.add_heading(category, level=2)
        for q, a in questions:
            p = doc.add_paragraph()
            run = p.add_run(f"T: {q}")
            run.bold = True

            p = doc.add_paragraph()
            p.add_run(f"J: {a}")
            doc.add_paragraph()

    doc.add_page_break()


def add_section_15(doc, images):
    """Section 15: Lampiran"""
    doc.add_heading("15. Lampiran", level=1)

    doc.add_heading("15.1 Contoh Data yang Dikirim ke Server", level=2)
    doc.add_paragraph("Berikut contoh format data yang dikirim gateway ke server (format JSON):")

    code = '''
{
  "device_id": "SRT-MGATE-001",
  "timestamp": "2025-11-30T10:30:45",
  "data": {
    "Suhu_Ruangan": {
      "value": 25.5,
      "unit": "°C"
    },
    "Kelembaban": {
      "value": 65,
      "unit": "%"
    }
  }
}
'''
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

    doc.add_heading("15.2 Daftar Port Umum", level=2)
    add_styled_table(doc,
        ["Layanan", "Port"],
        [
            ["MQTT (tanpa enkripsi)", "1883"],
            ["MQTT (dengan enkripsi)", "8883"],
            ["HTTP", "80"],
            ["HTTPS", "443"],
            ["Modbus TCP", "502"]
        ]
    )

    doc.add_heading("15.3 Daftar Baudrate Umum", level=2)
    add_styled_table(doc,
        ["Baudrate", "Penggunaan"],
        [
            ["9600", "Standar, paling umum"],
            ["19200", "Kecepatan menengah"],
            ["38400", "Kecepatan tinggi"],
            ["115200", "Kecepatan sangat tinggi"]
        ]
    )

    doc.add_page_break()


def add_section_16(doc, images):
    """Section 16: Kontak & Dukungan"""
    doc.add_heading("16. Kontak & Dukungan", level=1)

    doc.add_heading("Tim Dukungan SURIOTA", level=2)

    contact_info = [
        ("Email", "support@suriota.com"),
        ("Website", "www.suriota.com"),
        ("Jam Operasional", "Senin - Jumat, 08:00 - 17:00 WIB")
    ]

    for label, value in contact_info:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)

    doc.add_heading("Informasi Tambahan", level=2)
    doc.add_paragraph("Untuk dokumentasi teknis lebih lanjut, kunjungi:")
    links = [
        "GitHub Gateway: github.com/GifariKemal/GatewaySuriotaPOC",
        "GitHub Aplikasi: github.com/dickykhusnaedy/suriota_mobile_app"
    ]
    for link in links:
        doc.add_paragraph(link, style='List Bullet')


def add_credits_page(doc):
    """Add credits/author page at the end."""
    doc.add_page_break()

    # Add spacing
    for _ in range(6):
        doc.add_paragraph()

    # Document info
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Informasi Dokumen")
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = PRIMARY_COLOR

    doc.add_paragraph()

    info_table = doc.add_table(rows=6, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    info_data = [
        ("Judul", DOC_TITLE),
        ("Versi Dokumen", DOC_VERSION),
        ("Versi Aplikasi", APP_VERSION),
        ("Tanggal Pembuatan", DOC_DATE),
        ("Penulis", DOC_AUTHOR),
        ("Penerbit", DOC_COMPANY)
    ]

    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(row.cells[0], "F0F0F0")
        row.cells[0].width = Inches(2)
        row.cells[1].width = Inches(3.5)

    # Add spacing
    for _ in range(4):
        doc.add_paragraph()

    # Copyright
    copyright_para = doc.add_paragraph()
    copyright_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_para.add_run("© 2025 SURIOTA. Hak Cipta Dilindungi.")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # Final note
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("Dokumen ini dibuat untuk membantu pengguna dalam mengoperasikan "
                       "Gateway Config App dan SRT-MGATE-1210 Industrial IoT Gateway.")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def main():
    """Main function to generate the document."""
    print("=" * 60)
    print("Gateway Config App - User Guide Document Generator")
    print("=" * 60)

    # Create document
    print("\n[1/8] Creating document...")
    doc = Document()

    # Setup page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Setup styles
    print("[2/8] Setting up styles...")
    setup_styles(doc)

    # Get images
    print("[3/8] Loading images...")
    images = get_sorted_images()
    print(f"      Found {len(images)} images")

    # Add cover page
    print("[4/8] Creating cover page...")
    add_cover_page(doc)

    # Add table of contents
    print("[5/8] Creating table of contents...")
    add_table_of_contents(doc)

    # Add header and footer
    print("[6/8] Adding headers and footers...")
    add_header_footer(doc)

    # Add content sections
    print("[7/8] Adding content sections...")
    sections = [
        ("Section 1: Pendahuluan", add_section_1),
        ("Section 2: Persyaratan Sistem", add_section_2),
        ("Section 3: Instalasi Aplikasi", add_section_3),
        ("Section 4: Memulai Aplikasi", add_section_4),
        ("Section 5: Menghubungkan Gateway", add_section_5),
        ("Section 6: Menu Utama", add_section_6),
        ("Section 7: Perangkat Komunikasi", add_section_7),
        ("Section 8: Data Register", add_section_8),
        ("Section 9: Koneksi Server", add_section_9),
        ("Section 10: Status Perangkat", add_section_10),
        ("Section 11: Pengaturan & Backup", add_section_11),
        ("Section 12: Pengaturan Aplikasi", add_section_12),
        ("Section 13: Pemecahan Masalah", add_section_13),
        ("Section 14: FAQ", add_section_14),
        ("Section 15: Lampiran", add_section_15),
        ("Section 16: Kontak & Dukungan", add_section_16),
    ]

    for name, func in sections:
        print(f"      - {name}")
        func(doc, images)

    # Add credits page
    print("      - Credits Page")
    add_credits_page(doc)

    # Save document
    print(f"[8/8] Saving document to: {OUTPUT_FILE}")
    doc.save(OUTPUT_FILE)

    print("\n" + "=" * 60)
    print("Document generated successfully!")
    print(f"Output: {OUTPUT_FILE}")
    print("=" * 60)


if __name__ == "__main__":
    main()
